import streamlit as st
import os
import sys
from transformers import T5ForConditionalGeneration, T5Tokenizer
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS
from langchain.schema import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import torch

class T5Wrapper:
    def __init__(self, model_name, temperature=0.9, top_p=1, verbose=True, n_ctx=4096):
        self.device = torch.device("cpu")  # Ensure the model runs on CPU
        self.tokenizer = T5Tokenizer.from_pretrained(model_name)
        self.model = T5ForConditionalGeneration.from_pretrained(model_name).to(self.device)
        self.temperature = temperature
        self.top_p = top_p
        self.verbose = verbose
        self.n_ctx = n_ctx

    def generate(self, input_text):
        input_ids = self.tokenizer.encode(input_text, return_tensors="pt").to(self.device)
        outputs = self.model.generate(input_ids, do_sample=True, temperature=self.temperature, top_p=self.top_p, max_length=self.n_ctx)
        return self.tokenizer.decode(outputs[0], skip_special_tokens=True)

# Function to load documents from the directory
def load_documents_from_directory(directory_path):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                documents.append(Document(page_content=content))
        elif filename.endswith(".docx"):
            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            documents.append(Document(page_content=content))
        elif filename.endswith(".pdf"):
            pdf_reader = PdfReader(file_path)
            content = "\n".join([page.extract_text() for page in pdf_reader.pages])
            documents.append(Document(page_content=content))
    return documents

# Load documents from the specified directory
directory_path = "./data"  # Change this path as per your local directory structure
documents = load_documents_from_directory(directory_path)

# Create FAISS index from the documents
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vector_store = FAISS.from_documents(documents, embeddings)

# Initialize your T5 model
llm = T5Wrapper(model_name="google/flan-t5-base")

# Create a custom retrieval QA chain
class CustomRetrievalQA:
    def __init__(self, llm, vector_store):
        self.llm = llm
        self.vector_store = vector_store

    def run(self, query):
        results = self.vector_store.similarity_search(query, k=2)
        combined_text = " ".join([doc.page_content for doc in results])
        prompt = f"Context: {combined_text}\n\nQuestion: {query}\n\nAnswer:"
        return self.llm.generate(prompt)

# Initialize the QA system
qa = CustomRetrievalQA(llm=llm, vector_store=vector_store)

# Streamlit app
def main():
    st.title("Retrieval-based Question Answering System")

    user_input = st.text_input("Enter your query:")

    if st.button("Submit"):
        if user_input.strip() == '':
            st.error("Please enter a query.")
        else:
            result = qa.run(user_input)
            st.success(f"Answer: {result}")

if __name__ == "__main__":
    main()
